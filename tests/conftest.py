from io import BytesIO

import pytest
from docx import Document as DocxDocument

from src.models.sop import (
    SOPBranch,
    SOPDecision,
    SOPDocument,
    SOPElement,
    SOPElementType,
)


@pytest.fixture
def sample_sop_text() -> str:
    return (
        "Receive customer support email.\n"
        "Check if the issue is billing-related.\n"
        "If yes, assign to Billing Queue.\n"
        "If no, assign to General Support Queue.\n"
        "Send acknowledgment email to customer.\n"
        "Close the triage step."
    )


@pytest.fixture
def sample_sop_docx_bytes() -> bytes:
    """Create a minimal .docx in memory for testing."""
    doc = DocxDocument()
    doc.add_heading("Customer Support Triage SOP", level=1)
    doc.add_paragraph("1. Receive customer support email.")
    doc.add_paragraph("2. Check if the issue is billing-related.")
    doc.add_paragraph("   If yes, assign to Billing Queue.")
    doc.add_paragraph("   If no, assign to General Support Queue.")
    doc.add_paragraph("3. Send acknowledgment email to customer.")
    doc.add_paragraph("4. Close the triage step.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_sop_document() -> SOPDocument:
    """A pre-built SOPDocument matching the example SOP."""
    return SOPDocument(
        title="Customer Support Triage",
        elements=[
            SOPElement(element_type=SOPElementType.STEP, text="Receive customer support email"),
            SOPElement(
                element_type=SOPElementType.DECISION,
                text="Check if the issue is billing-related",
                decision=SOPDecision(
                    question="Is the issue billing-related?",
                    branches=[
                        SOPBranch(
                            condition_label="Yes",
                            steps=[
                                SOPElement(
                                    element_type=SOPElementType.STEP,
                                    text="Assign to Billing Queue",
                                )
                            ],
                        ),
                        SOPBranch(
                            condition_label="No",
                            steps=[
                                SOPElement(
                                    element_type=SOPElementType.STEP,
                                    text="Assign to General Support Queue",
                                )
                            ],
                        ),
                    ],
                ),
            ),
            SOPElement(
                element_type=SOPElementType.STEP,
                text="Send acknowledgment email to customer",
            ),
            SOPElement(element_type=SOPElementType.STEP, text="Close the triage step"),
        ],
    )


@pytest.fixture
def linear_sop_document() -> SOPDocument:
    """An SOP with no decisions — purely linear."""
    return SOPDocument(
        title="Simple Process",
        elements=[
            SOPElement(element_type=SOPElementType.STEP, text="Step A"),
            SOPElement(element_type=SOPElementType.STEP, text="Step B"),
            SOPElement(element_type=SOPElementType.STEP, text="Step C"),
        ],
    )
