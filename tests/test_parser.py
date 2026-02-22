from io import BytesIO

from docx import Document as DocxDocument

from src.parser.docx_parser import DocxSOPParser


class TestDocxTextExtraction:
    """Test the text extraction from .docx files (no LLM call)."""

    def test_extract_text_from_docx(self, sample_sop_docx_bytes):
        # We test the _extract_text method directly (no LLM)
        parser = DocxSOPParser.__new__(DocxSOPParser)
        text = parser._extract_text(sample_sop_docx_bytes)

        assert "Receive customer support email" in text
        assert "Check if the issue is billing-related" in text
        assert "If yes, assign to Billing Queue" in text
        assert "If no, assign to General Support Queue" in text
        assert "Send acknowledgment email" in text
        assert "Close the triage step" in text

    def test_extract_text_empty_doc(self):
        doc = DocxDocument()
        buf = BytesIO()
        doc.save(buf)

        parser = DocxSOPParser.__new__(DocxSOPParser)
        text = parser._extract_text(buf.getvalue())

        assert text == ""

    def test_extract_text_preserves_line_structure(self):
        doc = DocxDocument()
        doc.add_paragraph("First line")
        doc.add_paragraph("Second line")
        doc.add_paragraph("Third line")
        buf = BytesIO()
        doc.save(buf)

        parser = DocxSOPParser.__new__(DocxSOPParser)
        text = parser._extract_text(buf.getvalue())

        lines = text.split("\n")
        assert len(lines) == 3
        assert lines[0] == "First line"
        assert lines[1] == "Second line"
        assert lines[2] == "Third line"
